
import io
import json
import logging
import re
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import nbformat
import pandas as pd
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from openai import OpenAI, OpenAIError
from PyPDF2 import PdfReader

# ============================================================
# PDF-BASED KNOWLEDGE BASE RAG EVALUATOR
# 5 TASKS / 100 MARKS
#
# Task 1 = 15
# Task 2 = 15
# Task 3 = 20
# Task 4 = 25
# Task 5 = 25
# TOTAL   = 100
# ============================================================

st.set_page_config(
    page_title="PDF RAG Project Evaluator",
    page_icon="📊",
    layout="wide",
)

st.title("📊 PDF-Based Knowledge Base RAG — Evaluator")
st.caption(
    "Evaluates student submissions against the uploaded problem statement "
    "and the 5-task, 100-mark rubric."
)

OPENAI_MODEL = "gpt-4.1"
MAX_WORKERS = 4

# ============================================================
# OPENAI
# ============================================================

def get_openai_client() -> Optional[OpenAI]:
    api_key = st.secrets.get("OPENAI_API_KEY")
    if not api_key:
        st.error(
            "OPENAI_API_KEY is missing. Add it in Streamlit Secrets."
        )
        return None

    return OpenAI(
        api_key=api_key,
        timeout=180,
        max_retries=2,
    )


client = get_openai_client()

# ============================================================
# FILE READING
# ============================================================

def decode_bytes(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode(errors="ignore")


def read_pdf(raw: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(raw))
        parts = []

        for page_no, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(
                    f"[PDF PAGE {page_no}]\n{text.strip()}"
                )

        return "\n\n".join(parts)

    except Exception as exc:
        logging.exception("PDF parsing failed")
        return f"[PDF PARSE ERROR] {exc}"


def read_docx(raw: bytes) -> str:
    try:
        document = Document(io.BytesIO(raw))
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                parts.append(
                    " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                    )
                )

        return "\n".join(parts)

    except Exception as exc:
        logging.exception("DOCX parsing failed")
        return f"[DOCX PARSE ERROR] {exc}"


def read_html(raw: bytes) -> str:
    soup = BeautifulSoup(
        decode_bytes(raw),
        "html.parser",
    )

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    return soup.get_text(
        separator="\n",
        strip=True,
    )


def read_notebook(raw: bytes) -> str:
    try:
        notebook = nbformat.reads(
            decode_bytes(raw),
            as_version=4,
        )
    except Exception as exc:
        return f"[NOTEBOOK PARSE ERROR] {exc}"

    parts = []

    for index, cell in enumerate(
        notebook.cells,
        start=1,
    ):
        source = "".join(
            cell.get("source", "")
        ).strip()

        if not source:
            continue

        if cell.cell_type == "markdown":
            parts.append(
                f"[MARKDOWN CELL {index}]\n{source}"
            )

        elif cell.cell_type == "code":
            parts.append(
                f"[CODE CELL {index}]\n{source}"
            )

            # Include visible execution outputs where available.
            outputs = cell.get("outputs", [])

            for output in outputs:
                if "text" in output:
                    text = "".join(output["text"]).strip()
                    if text:
                        parts.append(
                            f"[OUTPUT CELL {index}]\n{text}"
                        )

                elif "data" in output:
                    data = output["data"]

                    if "text/plain" in data:
                        text = "".join(
                            data["text/plain"]
                        ).strip()

                        if text:
                            parts.append(
                                f"[OUTPUT CELL {index}]\n{text}"
                            )

    return "\n\n".join(parts)


def read_text(raw: bytes) -> str:
    return decode_bytes(raw)


# ============================================================
# SUBMISSION PARSING
# ============================================================

def empty_submission() -> Dict[str, Any]:
    return {
        "files": [],
        "documentation": [],
        "notebooks": [],
        "code": [],
        "other": [],
        "warnings": [],
    }


def add_file(
    result: Dict[str, Any],
    filename: str,
    raw: bytes,
) -> None:

    suffix = Path(filename).suffix.lower()
    result["files"].append(filename)

    try:

        if suffix == ".ipynb":
            result["notebooks"].append(
                f"FILE: {filename}\n"
                f"{read_notebook(raw)}"
            )

        elif suffix == ".py":
            result["code"].append(
                f"FILE: {filename}\n"
                f"{decode_bytes(raw)}"
            )

        elif suffix == ".pdf":
            result["documentation"].append(
                f"FILE: {filename}\n"
                f"{read_pdf(raw)}"
            )

        elif suffix == ".docx":
            result["documentation"].append(
                f"FILE: {filename}\n"
                f"{read_docx(raw)}"
            )

        elif suffix in {".html", ".htm"}:
            result["documentation"].append(
                f"FILE: {filename}\n"
                f"{read_html(raw)}"
            )

        elif suffix in {".md", ".txt"}:
            result["documentation"].append(
                f"FILE: {filename}\n"
                f"{read_text(raw)}"
            )

        elif suffix in {
            ".json",
            ".yaml",
            ".yml",
            ".sql",
            ".csv",
        }:
            result["other"].append(
                f"FILE: {filename}\n"
                f"{decode_bytes(raw)}"
            )

        else:
            result["warnings"].append(
                f"Unsupported file: {filename}"
            )

    except Exception as exc:
        logging.exception(
            "Unable to parse %s",
            filename,
        )
        result["warnings"].append(
            f"Unable to parse {filename}: {exc}"
        )


def parse_zip(raw: bytes) -> Dict[str, Any]:
    result = empty_submission()

    try:
        with zipfile.ZipFile(
            io.BytesIO(raw)
        ) as archive:

            for info in archive.infolist():

                if info.is_dir():
                    continue

                filename = info.filename

                # Ignore cache/hidden files.
                if (
                    "__pycache__" in filename
                    or filename.startswith(".")
                    or "/." in filename
                ):
                    continue

                add_file(
                    result,
                    filename,
                    archive.read(info),
                )

    except zipfile.BadZipFile:
        result["warnings"].append(
            "The uploaded ZIP is invalid."
        )

    return result


def parse_submission(uploaded_file) -> Dict[str, Any]:
    raw = uploaded_file.getvalue()
    suffix = Path(
        uploaded_file.name
    ).suffix.lower()

    if suffix == ".zip":
        return parse_zip(raw)

    result = empty_submission()

    add_file(
        result,
        uploaded_file.name,
        raw,
    )

    return result


def truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return (
        text[:limit]
        + "\n\n[CONTENT TRUNCATED FOR EVALUATION CONTEXT]"
    )


def build_submission_context(
    parsed: Dict[str, Any]
) -> str:

    documentation = "\n\n".join(
        parsed["documentation"]
    )

    notebooks = "\n\n".join(
        parsed["notebooks"]
    )

    code = "\n\n".join(
        parsed["code"]
    )

    other = "\n\n".join(
        parsed["other"]
    )

    return f"""
STUDENT SUBMISSION FILES
========================
{chr(10).join(parsed["files"])}

NOTEBOOK EVIDENCE
=================
{truncate(notebooks, 30000)}

PYTHON CODE EVIDENCE
====================
{truncate(code, 30000)}

DOCUMENTATION EVIDENCE
======================
{truncate(documentation, 20000)}

OTHER FILE EVIDENCE
===================
{truncate(other, 10000)}

PARSER WARNINGS
===============
{chr(10).join(parsed["warnings"])}
"""


# ============================================================
# PROBLEM STATEMENT
# ============================================================

def read_problem(uploaded_file) -> str:

    raw = uploaded_file.getvalue()
    suffix = Path(
        uploaded_file.name
    ).suffix.lower()

    if suffix == ".pdf":
        return read_pdf(raw)

    if suffix == ".docx":
        return read_docx(raw)

    if suffix in {".txt", ".md"}:
        return decode_bytes(raw)

    return ""


# ============================================================
# RUBRIC NORMALIZATION
# ============================================================

RUBRIC_ALIASES = {
    "criterion": [
        "criterion",
        "criteria",
        "evaluation criteria",
    ],
    "max_score": [
        "max score",
        "max marks",
        "marks",
        "score",
        "weight",
    ],
    "description": [
        "description",
        "evaluation parameters",
        "details",
        "evaluation description",
    ],
}


def normalize_column_name(
    value: Any
) -> str:

    value = str(value).strip().lower()

    value = value.replace(
        "_",
        " ",
    )

    return re.sub(
        r"\s+",
        " ",
        value,
    )


def find_matching_column(
    columns,
    aliases,
) -> Optional[str]:

    normalized = {
        normalize_column_name(column): column
        for column in columns
    }

    for alias in aliases:

        key = normalize_column_name(
            alias
        )

        if key in normalized:
            return normalized[key]

    return None


def normalize_rubric(
    df: pd.DataFrame
) -> Tuple[
    Optional[pd.DataFrame],
    List[str]
]:

    df = df.copy()

    df.columns = [
        str(column).strip()
        for column in df.columns
    ]

    mapping = {}
    missing = []

    for canonical_name, aliases in RUBRIC_ALIASES.items():

        column = find_matching_column(
            df.columns,
            aliases,
        )

        if column is None:
            missing.append(
                canonical_name
            )
        else:
            mapping[column] = canonical_name

    if missing:
        return None, missing

    normalized = df.rename(
        columns=mapping
    )

    normalized = normalized[
        [
            "criterion",
            "max_score",
            "description",
        ]
    ]

    normalized = normalized.dropna(
        subset=[
            "criterion",
            "max_score",
        ],
        how="any",
    )

    normalized["criterion"] = (
        normalized["criterion"]
        .astype(str)
        .str.strip()
    )

    normalized["description"] = (
        normalized["description"]
        .fillna("")
        .astype(str)
        .str.strip()
    )

    normalized["max_score"] = pd.to_numeric(
        normalized["max_score"],
        errors="coerce",
    )

    normalized = normalized.dropna(
        subset=["max_score"]
    )

    return normalized, []


def rubric_to_text(
    df: pd.DataFrame
) -> str:

    lines = []

    for _, row in df.iterrows():

        lines.append(
            f"""
Criterion:
{row["criterion"]}

Maximum Score:
{row["max_score"]}

Evaluation Description:
{row["description"]}
"""
        )

    return "\n".join(lines)


# ============================================================
# EVALUATION PROMPT
# ============================================================

def build_evaluation_prompt(
    problem_text: str,
    rubric_text: str,
    submission_context: str,
    custom_instructions: str,
) -> str:

    return f"""
You are a fair, implementation-tolerant evaluator for a guided
PDF-Based Knowledge Base RAG project.

Use ONLY:
1. The supplied problem statement.
2. The supplied evaluation rubric.
3. Evidence found in the student's submitted files.

============================================================
PROBLEM STATEMENT
============================================================

{problem_text}

============================================================
RUBRIC
============================================================

{rubric_text}

============================================================
STUDENT SUBMISSION EVIDENCE
============================================================

{submission_context}

============================================================
SCORING PHILOSOPHY — VERY IMPORTANT
============================================================

This is a GUIDED PROJECT.

Be fair and implementation-tolerant.

If the student's implementation APPROXIMATELY satisfies the task
requirement and clearly demonstrates the intended functionality,
AWARD FULL MARKS for that task.

Do NOT reduce marks for:

- Different but valid implementation approaches.
- Different variable names.
- Different function names.
- Different code organization.
- Different UI design.
- Different but valid LangChain patterns.
- Minor configuration differences that do not materially affect
  the required functionality.
- Minor documentation differences.
- Reasonable alternative implementation patterns.
- Small deviations that do not affect the intended outcome.

Judge FUNCTIONALITY and REQUIREMENT SATISFACTION, not code similarity
to a reference solution.

FULL-SCORE RULE:
If the task's core requirement is implemented correctly or
approximately correctly and there is no major missing component,
award the FULL MAXIMUM SCORE for that task.

Do NOT subtract 1–5 marks simply because of minor imperfections.
Do NOT create artificial partial scores.

Examples:
- Valid alternative implementation achieving the same task outcome = FULL.
- Minor chunk overlap difference while fixed-size chunking is demonstrated = FULL.
- Different valid retriever configuration = FULL.
- Different prompt wording with proper grounding = FULL.
- Different UI layout while upload, QA and sources work = FULL.
- README is brief but adequate = FULL when the implementation is otherwise complete.

Only reduce marks when there is a MEANINGFUL deficiency, such as:

- A required component is missing.
- A required technology is clearly not used.
- A major requirement is incorrectly implemented.
- The functionality is substantially incomplete.
- The implementation is clearly nonfunctional.
- There is no meaningful evidence of implementation.

Do NOT penalize minor imperfections.

Do NOT require the student to reproduce the exact reference code.

A student's code can be different from the solution and still receive
full marks.

============================================================
IMPORTANT TECHNICAL CHECKS
============================================================

Check the student's evidence for:

Task 1:
- PDF input/provision.
- LangChain PDF loader.
- Successful document extraction.
- Page/source metadata where applicable.

Task 2:
- Fixed-size text chunking.
- Appropriate text splitter.
- Chunk size around/at 1000.
- Configurable chunk overlap.
- Generated chunks.

Task 3:
- text-embedding-3-small.
- Embedding generation.
- ChromaDB.
- Document chunks/embeddings/metadata stored.
- Persistence/reuse where applicable.

Task 4:
- Retriever.
- Similarity search.
- Relevant Top-K chunks.
- Context construction.
- GPT-4o-mini.
- Answer based on retrieved context.
- Reasonable grounding / no unsupported hallucination.

Task 5:
- Final application.
- User PDF upload/provision.
- Complete end-to-end pipeline.
- User question interface.
- Answer display.
- Source/page display where applicable.
- Handling of questions outside the document.
- Code/documentation/README.

============================================================
SCORING RULES
============================================================

There are exactly 5 tasks.

Task 1 maximum = 15
Task 2 maximum = 15
Task 3 maximum = 20
Task 4 maximum = 25
Task 5 maximum = 25

TOTAL MAXIMUM = 100.

The final score MUST NEVER exceed 100.

If a task is substantially satisfied, award its full task score.

Examples:

- Correct RAG implementation using a slightly different valid splitter:
  FULL MARKS.

- Correct ChromaDB implementation with different persistence setup:
  FULL MARKS.

- Correct retrieval using a valid LangChain retriever pattern:
  FULL MARKS.

- Streamlit interface looks different from reference:
  FULL MARKS if required functionality works.

- Code differs from reference solution but achieves the requirement:
  FULL MARKS.

Partial marks should be used only when a meaningful part of a task
is actually missing or incorrect.

============================================================
EVIDENCE RULE
============================================================

Award marks based on evidence.

Code is evidence.
Notebook cells are evidence.
Notebook outputs are evidence.
Documentation is evidence.
README is evidence.

A claim in a README alone is not strong evidence that an implementation
works if corresponding code is absent.

However, do not require execution output when the submitted code clearly
and correctly implements the requirement.

============================================================
TASK-LEVEL FULL-SCORE GUIDANCE
============================================================

TASK 1 / 15:
Give 15/15 when the student loads/provides a PDF using a valid
LangChain PDF loader and extracts usable document content.
Minor metadata or path differences should NOT reduce the score.
Reduce only if PDF loading/document processing is materially missing
or nonfunctional.

TASK 2 / 15:
Give 15/15 when fixed-size chunking is implemented and the student
demonstrates an appropriate chunk size and configurable overlap.
A small difference in overlap or equivalent valid splitter should
NOT reduce the score. Reduce only for materially missing/incorrect
chunking.

TASK 3 / 20:
Give 20/20 when embeddings are generated with the required
text-embedding-3-small model (or the implementation clearly shows
the intended required embedding model) and stored in ChromaDB.
Minor persistence/configuration differences should NOT reduce the score.

TASK 4 / 25:
Give 25/25 when retrieval works, relevant chunks are used as context,
GPT-4o-mini is integrated, and the answer is grounded in retrieved
context. Valid alternative retriever/prompt patterns receive FULL.
Reduce only when a core retrieval/LLM/grounding requirement is missing
or materially incorrect.

TASK 5 / 25:
Give 25/25 when the final application demonstrates the complete
working PDF RAG flow, including PDF input/upload/provision, user
questioning, answer generation, and appropriate source/context
handling. UI differences and minor documentation gaps should NOT
reduce the score if the intended application functionality is present.

IMPORTANT:
A student who completes the core functionality for all five tasks
should normally receive 100/100 even when the implementation differs
from the reference solution.

============================================================
OUTPUT
============================================================

Return ONLY valid JSON.

Do NOT return numeric task scores.

Classify each task using exactly one status:
- MET
- PARTIALLY_MET
- NOT_MET

If the core task requirement is correctly or approximately satisfied,
return MET. Minor implementation differences must not cause a partial status.

Use PARTIALLY_MET only for a meaningful missing or materially incorrect
part of the task.

Use NOT_MET only when the task is essentially missing or nonfunctional.

Required JSON:

{
  "task_status": {
    "Task 1 – PDF Loading & Document Processing": "MET",
    "Task 2 – Text Chunking": "MET",
    "Task 3 – Embeddings & ChromaDB": "MET",
    "Task 4 – Retrieval & RAG Answer Generation": "MET",
    "Task 5 – Final PDF RAG Application": "MET"
  },
  "criterion_feedback": {
    "Task 1 – PDF Loading & Document Processing": "...",
    "Task 2 – Text Chunking": "...",
    "Task 3 – Embeddings & ChromaDB": "...",
    "Task 4 – Retrieval & RAG Answer Generation": "...",
    "Task 5 – Final PDF RAG Application": "..."
  },
  "area_of_strength": "...",
  "area_of_improvement": "...",
  "overall_evaluation": "...",
  "evidence_summary": [],
  "missing_requirements": []
}

Do not return a numeric score. Python will calculate the score.

IMPORTANT:
Area of Improvement is advisory only.
Do not lower a task status because of:
- UI polish,
- visual design,
- code style,
- minor documentation gaps,
- optional enhancements,
- production-readiness suggestions.

Those items may be mentioned in Area of Improvement while the student
still receives full marks.


============================================================
ADDITIONAL EVALUATOR INSTRUCTIONS
============================================================

{custom_instructions}
"""


# ============================================================
# JSON / SCORE HANDLING
# ============================================================

def parse_json(
    raw: str
) -> Dict[str, Any]:

    try:
        data = json.loads(raw)

    except json.JSONDecodeError:

        match = re.search(
            r"\{.*\}",
            raw or "",
            re.DOTALL,
        )

        if not match:
            return {
                "task_status": {},
                "criterion_feedback": {},
                "area_of_strength": "",
                "area_of_improvement": "",
                "overall_evaluation": (
                    "The evaluator model did not return valid JSON."
                ),
                "evidence_summary": [],
                "missing_requirements": [],
            }

        try:
            data = json.loads(
                match.group()
            )

        except json.JSONDecodeError:

            return {
                "task_status": {},
                "criterion_feedback": {},
                "area_of_strength": "",
                "area_of_improvement": "",
                "overall_evaluation": (
                    "The evaluator model returned invalid JSON."
                ),
                "evidence_summary": [],
                "missing_requirements": [],
            }

    if not isinstance(data, dict):
        data = {}

    data.setdefault(
        "task_status",
        {},
    )

    # Backward compatibility: if a model returns numeric "scores"
    # instead of statuses, convert them to statuses rather than giving 0.
    if not data.get("task_status") and isinstance(data.get("scores"), dict):
        data["task_status"] = {
            str(k): (
                "MET"
                if isinstance(v, (int, float)) and float(v) > 0
                else str(v)
            )
            for k, v in data["scores"].items()
        }

    data.setdefault(
        "criterion_feedback",
        {},
    )

    data.setdefault(
        "area_of_strength",
        "",
    )

    data.setdefault(
        "area_of_improvement",
        "",
    )

    data.setdefault(
        "overall_evaluation",
        "",
    )

    data.setdefault(
        "evidence_summary",
        [],
    )

    data.setdefault(
        "missing_requirements",
        [],
    )

    return data


def coerce_score(
    value: Any,
    maximum: float,
) -> float:

    try:
        score = float(value)
    except (
        TypeError,
        ValueError,
    ):
        return 0.0

    return max(
        0.0,
        min(
            score,
            float(maximum),
        ),
    )



def infer_task_status_from_evidence(
    parsed_submission: Dict[str, Any],
    rubric_df: pd.DataFrame,
) -> Dict[str, str]:
    """
    Guided-project fallback.
    If the LLM response has no usable task_status, infer MET from clear
    implementation evidence in the student's submitted notebook/code/html.
    This prevents a malformed LLM JSON response from turning a completed
    guided project into 0.
    """
    combined = " ".join([
        " ".join(parsed_submission.get("notebooks", [])),
        " ".join(parsed_submission.get("code", [])),
        " ".join(parsed_submission.get("documentation", [])),
        " ".join(parsed_submission.get("other", [])),
    ]).lower()

    def has_any(*terms):
        return any(term.lower() in combined for term in terms)

    statuses = {}

    for _, row in rubric_df.iterrows():
        criterion = str(row["criterion"]).strip()

        if criterion.lower().startswith("task 1"):
            met = (
                has_any("pypdfloader", "pdf loader", "pdf document loader")
                and has_any("load()", ".load(", "documents =")
            )
        elif criterion.lower().startswith("task 2"):
            met = (
                has_any(
                    "recursivecharactertextsplitter",
                    "text splitter",
                    "character text splitter",
                )
                and has_any("chunk_size", "chunk size")
                and has_any("chunk_overlap", "chunk overlap", "overlap")
            )
        elif criterion.lower().startswith("task 3"):
            met = (
                has_any("text-embedding-3-small")
                and has_any("chromadb", "from langchain_chroma import chroma", "chroma(")
                and has_any(
                    "vectorstore",
                    "vector store",
                    "from_documents",
                    "similarity_search",
                )
            )
        elif criterion.lower().startswith("task 4"):
            met = (
                has_any("retriever", "as_retriever", "similarity_search")
                and has_any("gpt-4o-mini")
                and has_any("context", "retrieved")
                and has_any("llm.invoke", "invoke(")
            )
        elif criterion.lower().startswith("task 5"):
            met = (
                has_any(
                    "st.file_uploader",
                    "file_uploader",
                    "upload pdf",
                    "upload_pdf",
                    "uploaded_file",
                )
                and has_any(
                    "streamlit",
                    "st.button",
                    "ask a question",
                    "input(",
                    "question =",
                )
                and has_any(
                    "gpt-4o-mini",
                    "retriever",
                    "vectorstore",
                )
            )
        else:
            met = False

        statuses[criterion] = "MET" if met else "NOT_MET"

    return statuses


def get_score_from_status(
    task_status: Dict[str, Any],
    criterion: str,
    maximum: float,
) -> float:

    status = task_status.get(criterion)

    if status is None:
        normalized = {
            normalize_column_name(k): v
            for k, v in task_status.items()
        }
        status = normalized.get(
            normalize_column_name(criterion),
            "NOT_MET",
        )

    if isinstance(status, bool):
        return float(maximum) if status else 0.0

    # Backward compatibility for a numeric model response.
    if isinstance(status, (int, float)):
        numeric = float(status)
        if numeric > 0:
            # In a guided project, a positive score from the model means
            # the task was considered substantially completed. Give full marks.
            return float(maximum)
        return 0.0

    status = str(status).strip().upper().replace("-", "_").replace(" ", "_")

    if status in {"MET", "COMPLETE", "COMPLETED", "SATISFIED", "PASS", "FULL"}:
        return float(maximum)

    if status in {"PARTIALLY_MET", "PARTIAL", "PARTIALLY_COMPLETE", "INCOMPLETE"}:
        return round(float(maximum) * 0.50, 2)

    return 0.0


def list_to_text(
    value: Any
) -> str:

    if isinstance(
        value,
        list,
    ):
        return "; ".join(
            str(item)
            for item in value
            if str(item).strip()
        )

    return str(
        value or ""
    )


# ============================================================
# LLM EVALUATION
# ============================================================

def evaluate_submission(
    prompt: str,
) -> Dict[str, Any]:

    if client is None:
        raise RuntimeError(
            "OpenAI client is not configured."
        )

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        temperature=0,
        response_format={
            "type": "json_object"
        },
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a strict but fair academic evaluator. "
                    "Use the supplied rubric as the scoring authority. "
                    "Be implementation-tolerant. "
                    "Return MET whenever the core requirement is satisfied "
                    "or approximately satisfied. "
                    "Do not invent partial deductions for minor differences."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
    )

    return parse_json(
        response.choices[0].message.content
        or "{}"
    )


# ============================================================
# ONE STUDENT
# ============================================================

def evaluate_one_student(
    uploaded_file,
    problem_text: str,
    rubric_df: pd.DataFrame,
    rubric_text: str,
    custom_instructions: str,
) -> Dict[str, Any]:

    parsed = parse_submission(
        uploaded_file
    )

    submission_context = (
        build_submission_context(
            parsed
        )
    )

    has_evidence = any(
        parsed[key]
        for key in [
            "documentation",
            "notebooks",
            "code",
            "other",
        ]
    )

    if not has_evidence:

        evaluation = {
            "task_status": {},
            "criterion_feedback": {},
            "area_of_strength": "",
            "area_of_improvement": (
                "No readable implementation evidence was found."
            ),
            "overall_evaluation": (
                "Submission could not be evaluated because "
                "no readable project evidence was found."
            ),
            "evidence_summary": [],
            "missing_requirements": [],
        }

    else:

        prompt = build_evaluation_prompt(
            problem_text=problem_text,
            rubric_text=rubric_text,
            submission_context=submission_context,
            custom_instructions=custom_instructions,
        )

        try:
            evaluation = evaluate_submission(
                prompt
            )

        except (
            OpenAIError,
            RuntimeError,
            TimeoutError,
            ValueError,
        ) as exc:

            logging.exception(
                "Evaluation failed for %s",
                uploaded_file.name,
            )

            evaluation = {
                "task_status": {},
                "criterion_feedback": {},
                "area_of_strength": "",
                "area_of_improvement": str(
                    exc
                ),
                "overall_evaluation": (
                    "Evaluation failed."
                ),
                "evidence_summary": [],
                "missing_requirements": [],
            }

    task_status = evaluation.get(
        "task_status",
        {},
    )

    # If the model did not return usable task statuses, infer them from
    # direct submission evidence instead of defaulting the project to zero.
    if not isinstance(task_status, dict) or not task_status:
        task_status = infer_task_status_from_evidence(
            parsed,
            rubric_df,
        )

    evidence_status = infer_task_status_from_evidence(
        parsed,
        rubric_df,
    )

    if isinstance(task_status, dict):
        for criterion_name, inferred_status in evidence_status.items():
            current_status = task_status.get(criterion_name)
            if current_status in (None, "", "NOT_MET", "UNKNOWN"):
                task_status[criterion_name] = inferred_status

    feedback = evaluation.get(
        "criterion_feedback",
        {},
    )

    result = {
        "Participant": uploaded_file.name,
    }

    total = 0.0
    maximum_total = 0.0

    # Only rubric rows determine score.
    for _, row in rubric_df.iterrows():

        criterion = str(
            row["criterion"]
        ).strip()

        maximum = float(
            row["max_score"]
        )

        score = get_score_from_status(
            task_status,
            criterion,
            maximum,
        )

        result[
            criterion
        ] = score

        result[
            f"{criterion} Feedback"
        ] = str(
            feedback.get(
                criterion,
                feedback.get(
                    normalize_column_name(
                        criterion
                    ),
                    "",
                ),
            )
        )

        total += score
        maximum_total += maximum

    # Deterministic full-score rule:
    # if all five tasks are MET, the final score is exactly 100/100.
    all_tasks_met = True

    for _, rubric_row in rubric_df.iterrows():
        criterion_name = str(rubric_row["criterion"]).strip()
        status = task_status.get(criterion_name)

        if status is None:
            normalized_status = {
                normalize_column_name(k): str(v).strip().upper()
                for k, v in task_status.items()
            }
            status = normalized_status.get(
                normalize_column_name(criterion_name),
                "NOT_MET",
            )

        status_norm = str(status).strip().upper().replace("-", "_").replace(" ", "_")
        if status_norm not in {"MET", "COMPLETE", "COMPLETED", "SATISFIED", "PASS", "FULL"}:
            all_tasks_met = False
            break

    if all_tasks_met and len(rubric_df) == 5:
        total = 100.0

    # Enforce maximum of 100.
    total = min(
        total,
        100.0,
    )

    result["Total Score"] = round(
        total,
        2,
    )

    result["Max Score"] = 100

    result["Percentage"] = round(
        total,
        2,
    )

    if total >= 90:
        rating = "Excellent"

    elif total >= 75:
        rating = "Good"

    elif total >= 60:
        rating = "Satisfactory"

    else:
        rating = "Needs Improvement"

    result["Rating"] = rating

    result[
        "Area of Strength"
    ] = str(
        evaluation.get(
            "area_of_strength",
            "",
        )
    )

    result[
        "Area of Improvement"
    ] = str(
        evaluation.get(
            "area_of_improvement",
            "",
        )
    )

    result[
        "Overall Evaluation"
    ] = str(
        evaluation.get(
            "overall_evaluation",
            "",
        )
    )

    # Combine all narrative feedback into ONE result field.
    task_feedback_parts = []
    for _, rubric_row in rubric_df.iterrows():
        criterion = str(rubric_row["criterion"]).strip()
        feedback_text = str(
            result.get(f"{criterion} Feedback", "")
        ).strip()
        if feedback_text:
            task_feedback_parts.append(
                f"{criterion}: {feedback_text}"
            )

    result["Feedback"] = "\n".join(task_feedback_parts)

    result[
        "Evidence Summary"
    ] = list_to_text(
        evaluation.get(
            "evidence_summary",
            [],
        )
    )

    result[
        "Missing Requirements"
    ] = list_to_text(
        evaluation.get(
            "missing_requirements",
            [],
        )
    )

    result[
        "Parser Warnings"
    ] = list_to_text(
        parsed.get(
            "warnings",
            [],
        )
    )

    return result


# ============================================================
# UI
# ============================================================

st.sidebar.header(
    "Evaluation Inputs"
)

problem_file = st.sidebar.file_uploader(
    "Problem Statement",
    type=[
        "pdf",
        "docx",
        "txt",
        "md",
    ],
)

rubric_file = st.sidebar.file_uploader(
    "Evaluation Rubric",
    type=["xlsx"],
)

submission_files = st.sidebar.file_uploader(
    "Student Submission(s)",
    type=[
        "ipynb",
        "zip",
        "docx",
        "pdf",
        "html",
        "htm",
        "py",
        "md",
        "txt",
    ],
    accept_multiple_files=True,
)

custom_instructions = st.text_area(
    "Additional Evaluation Instructions",
    value=(
        "Give full marks when the student approximately satisfies "
        "the requirement and demonstrates the intended functionality. "
        "Do not penalize minor implementation differences. "
        "Evaluate functionality rather than similarity to a reference solution."
    ),
    height=120,
)

# ============================================================
# RUN
# ============================================================

if st.button(
    "🚀 Evaluate Student Submission(s)",
    type="primary",
):

    if client is None:
        st.error(
            "OpenAI API configuration is missing."
        )
        st.stop()

    if problem_file is None:
        st.error(
            "Upload the problem statement."
        )
        st.stop()

    if rubric_file is None:
        st.error(
            "Upload the evaluation rubric."
        )
        st.stop()

    if not submission_files:
        st.error(
            "Upload at least one student submission."
        )
        st.stop()

    # --------------------------------------------------------
    # Problem
    # --------------------------------------------------------

    problem_text = read_problem(
        problem_file
    )

    if not problem_text.strip():
        st.warning(
            "Problem statement text could not be extracted."
        )

    # --------------------------------------------------------
    # Rubric
    # --------------------------------------------------------

    try:
        raw_rubric = pd.read_excel(
            rubric_file
        )

    except Exception as exc:

        st.error(
            f"Unable to read rubric: {exc}"
        )
        st.stop()

    rubric_df, missing = normalize_rubric(
        raw_rubric
    )

    if rubric_df is None:

        st.error(
            "Rubric must contain: "
            "Criterion, Max Score, Description."
        )

        st.write(
            "Columns detected:",
            list(raw_rubric.columns),
        )

        st.stop()

    # --------------------------------------------------------
    # Verify total = 100
    # --------------------------------------------------------

    if len(rubric_df) != 5:
        st.error(
            f"Evaluator expects exactly 5 task rows, but {len(rubric_df)} "
            "task rows were detected after removing summary rows."
        )
        st.dataframe(
            rubric_df[["criterion", "max_score", "description"]],
            use_container_width=True
        )
        st.stop()

    rubric_total = float(
        rubric_df["max_score"].sum()
    )

    if abs(
        rubric_total - 100
    ) > 0.001:

        st.error(
            f"Rubric total is {rubric_total}, "
            "but this evaluator requires exactly 100 marks."
        )
        st.stop()

    rubric_text = rubric_to_text(
        rubric_df
    )

    with st.expander(
        "Rubric Used",
        expanded=False,
    ):
        st.dataframe(
            rubric_df,
            use_container_width=True,
        )

        st.write(
            f"Rubric Total: {rubric_total}/100"
        )

    # --------------------------------------------------------
    # Evaluate
    # --------------------------------------------------------

    results = []

    progress = st.progress(
        0
    )

    status = st.empty()

    with ThreadPoolExecutor(
        max_workers=min(
            MAX_WORKERS,
            len(submission_files),
        )
    ) as executor:

        futures = {
            executor.submit(
                evaluate_one_student,
                file,
                problem_text,
                rubric_df,
                rubric_text,
                custom_instructions,
            ): file.name
            for file in submission_files
        }

        for index, future in enumerate(
            as_completed(futures),
            start=1,
        ):

            filename = futures[
                future
            ]

            try:

                result = future.result()
                results.append(
                    result
                )

            except Exception as exc:

                logging.exception(
                    "Unexpected error"
                )

                results.append(
                    {
                        "Participant": filename,
                        "Task 1 – PDF Loading & Document Processing": 0,
                        "Task 2 – Text Chunking": 0,
                        "Task 3 – Embeddings & ChromaDB": 0,
                        "Task 4 – Retrieval & RAG Answer Generation": 0,
                        "Task 5 – Final PDF RAG Application": 0,
                        "Total Score": 0,
                        "Max Score": 100,
                        "Percentage": 0,
                        "Rating": "Evaluation Failed",
                        "Area of Strength": "",
                        "Area of Improvement": str(
                            exc
                        ),
                        "Overall Evaluation": (
                            "Evaluation failed."
                        ),
                    }
                )

            status.info(
                f"Evaluated {index}/"
                f"{len(futures)}: {filename}"
            )

            progress.progress(
                index / len(futures)
            )

    if not results:
        st.warning(
            "No evaluation results were produced."
        )
        st.stop()

    results_df = pd.DataFrame(
        results
    )

    # ========================================================
    # FINAL SUMMARY
    # ========================================================

    st.subheader(
        "📊 Final Evaluation Summary"
    )

    summary_columns = [
        "Participant",
        "Task 1 – PDF Loading & Document Processing",
        "Task 2 – Text Chunking",
        "Task 3 – Embeddings & ChromaDB",
        "Task 4 – Retrieval & RAG Answer Generation",
        "Task 5 – Final PDF RAG Application",
        "Total Score",
        "Rating",
    ]

    summary_columns = [
        col
        for col in summary_columns
        if col in results_df.columns
    ]

    st.dataframe(
        results_df[
            summary_columns
        ],
        use_container_width=True,
    )

    # ========================================================
    # INDIVIDUAL STUDENT DETAILS
    # ========================================================

    st.subheader(
        "📝 Detailed Evaluation"
    )

    for _, student in results_df.iterrows():

        participant = student.get(
            "Participant",
            "Student",
        )

        total = student.get(
            "Total Score",
            0,
        )

        rating = student.get(
            "Rating",
            "",
        )

        with st.expander(
            f"{participant} — "
            f"{total}/100 — {rating}"
        ):

            st.markdown(
                "### Task-wise Score"
            )

            task_rows = []

            for _, rubric_row in rubric_df.iterrows():

                criterion = str(
                    rubric_row["criterion"]
                )

                task_rows.append(
                    {
                        "Task": criterion,
                        "Score": student.get(
                            criterion,
                            0,
                        ),
                        "Maximum": rubric_row[
                            "max_score"
                        ],
                        "Feedback": student.get(
                            f"{criterion} Feedback",
                            "",
                        ),
                    }
                )

            st.dataframe(
                pd.DataFrame(
                    task_rows
                ),
                use_container_width=True,
            )

            st.markdown(
                f"### Total Score: "
                f"**{total}/100**"
            )

            st.markdown(
                "### Area of Strength"
            )

            st.write(
                student.get(
                    "Area of Strength",
                    "",
                )
            )

            st.markdown(
                "### Area of Improvement"
            )

            st.write(
                student.get(
                    "Area of Improvement",
                    "",
                )
            )

            st.markdown(
                "### Overall Evaluation"
            )

            st.write(
                student.get(
                    "Overall Evaluation",
                    "",
                )
            )

            if student.get(
                "Missing Requirements",
                "",
            ):
                st.markdown(
                    "### Missing Requirements"
                )

                st.write(
                    student.get(
                        "Missing Requirements",
                        "",
                    )
                )

    # ========================================================
    # DOWNLOAD EXCEL
    # ========================================================

    excel_buffer = io.BytesIO()

    with pd.ExcelWriter(
        excel_buffer,
        engine="xlsxwriter",
    ) as writer:

        # Sheet 1: final student-level evaluation report.
        # The narrative fields are deliberately placed in the main result file.
        final_columns = [
            "Participant",
            "Task 1 – PDF Loading & Document Processing",
            "Task 2 – Text Chunking",
            "Task 3 – Embeddings & ChromaDB",
            "Task 4 – Retrieval & RAG Answer Generation",
            "Task 5 – Final PDF RAG Application",
            "Total Score",
            "Rating",
            "Area of Strength",
            "Area of Improvement",
            "Overall Evaluation",
            "Feedback",
        ]

        final_columns = [
            col
            for col in final_columns
            if col in results_df.columns
        ]

        results_df[
            final_columns
        ].to_excel(
            writer,
            index=False,
            sheet_name="Final Evaluation",
        )

        # Sheet 2: task-wise feedback
        detail_rows = []

        for _, student in results_df.iterrows():

            for _, rubric_row in rubric_df.iterrows():

                criterion = str(
                    rubric_row["criterion"]
                )

                detail_rows.append(
                    {
                        "Participant": student.get(
                            "Participant",
                            "",
                        ),
                        "Task": criterion,
                        "Score": student.get(
                            criterion,
                            0,
                        ),
                        "Maximum Score": rubric_row[
                            "max_score"
                        ],
                        "Task Feedback": student.get(
                            f"{criterion} Feedback",
                            "",
                        ),
                    }
                )

        pd.DataFrame(
            detail_rows
        ).to_excel(
            writer,
            index=False,
            sheet_name="Task-wise Details",
        )

    st.markdown("### Final Performance Summary")
    summary_display = results_df[
        [c for c in [
            "Participant",
            "Total Score",
            "Rating",
            "Area of Strength",
            "Area of Improvement",
            "Overall Evaluation",
            "Feedback",
        ] if c in results_df.columns]
    ]
    st.dataframe(
        summary_display,
        use_container_width=True,
    )

    st.download_button(
        "📥 Download Final Evaluation Excel",
        data=excel_buffer.getvalue(),
        file_name=(
            "PDF_RAG_Student_Evaluation_Report.xlsx"
        ),
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )

    st.success(
        "Evaluation completed. "
        "Final score is strictly out of 100."
    )
